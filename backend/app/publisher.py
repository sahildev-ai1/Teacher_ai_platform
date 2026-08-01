"""
Stage 10: Publishing.

Assembles the final TeacherKnowledgePackage and renders it into every format a
teacher might want: the canonical JSON, a Markdown "Teacher Guide", HTML (for
in-browser preview / copy button), a PDF (via xhtml2pdf - pure Python, no
system Cairo/Pango dependency), and a DOCX (via python-docx, walking the same
Markdown so there's exactly one source of truth for content).
"""
import io
import re

import markdown as md_lib
from docx import Document as DocxDocument
from docx.shared import Pt
from xhtml2pdf import pisa

from .schemas import TeacherKnowledgePackage


def build_package(
    document_id, classification, knowledge_base, teaching_plan,
    period_content, activities, assessments, learning_gaps, validation,
) -> TeacherKnowledgePackage:
    return TeacherKnowledgePackage(
        document_id=document_id,
        classification=classification,
        knowledge_base=knowledge_base,
        teaching_plan=teaching_plan,
        period_content=period_content,
        activities=activities,
        assessments=assessments,
        learning_gaps=learning_gaps,
        validation=validation,
    )


def render_markdown(pkg: TeacherKnowledgePackage) -> str:
    c, kb, plan = pkg.classification, pkg.knowledge_base, pkg.teaching_plan
    lines = []
    lines.append(f"# Teacher Guide: {c.topic}")
    lines.append(f"**Subject:** {c.subject} | **Grade:** {c.grade} | "
                  f"**Difficulty:** {c.difficulty} | **Chapter:** {c.chapter} | "
                  f"**Category:** {c.category} | **Language:** {c.language}\n")

    lines.append("## Learning Objectives")
    lines += [f"- {o}" for o in kb.learning_objectives]

    lines.append("\n## Prerequisites")
    lines += [f"- {p}" for p in kb.prerequisites]

    lines.append("\n## Key Concepts")
    for concept in kb.concepts:
        lines.append(f"**{concept.name}** — {concept.explanation}")

    lines.append("\n## Definitions")
    for d in kb.definitions:
        lines.append(f"- **{d.name}**: {d.explanation}")

    if kb.formulae:
        lines.append("\n## Formulae")
        lines += [f"- `{f}`" for f in kb.formulae]

    lines.append("\n## Common Misconceptions")
    for m in kb.common_misconceptions:
        lines.append(f"- ❌ {m.misconception} → ✅ {m.correction}")

    lines.append(f"\n## Teaching Plan ({plan.total_periods} periods)")
    lines.append(f"*{plan.rationale}*")

    activities_by_period = {a.period_number: a for a in pkg.activities}
    assessments_by_period = {a.period_number: a for a in pkg.assessments}
    content_by_period = {pc.period_number: pc for pc in pkg.period_content}

    for p in plan.periods:
        lines.append(f"\n---\n### Period {p.period_number}: {p.title} ({p.duration_minutes} min)")
        lines.append(f"**Objectives:** {', '.join(p.objectives)}")
        lines.append(f"**Pacing notes:** {p.pacing_notes}")

        pc = content_by_period.get(p.period_number)
        if pc:
            lines.append(f"\n**Entry Ticket:** {pc.entry_ticket}")
            lines.append(f"\n**Teacher Script:**\n{pc.teacher_script}")
            lines.append(f"\n**Blackboard Notes:**\n{pc.blackboard_notes}")
            lines.append(f"\n**Checkpoint Questions:**")
            lines += [f"- {q}" for q in pc.checkpoint_questions]
            lines.append(f"\n**Exit Ticket:** {pc.exit_ticket}")
            lines.append(f"\n**Homework:** {pc.homework}")
            lines.append(f"\n**Mentor Moment:** {pc.mentor_moment}")

        act = activities_by_period.get(p.period_number)
        if act:
            lines.append(f"\n**Activity — {act.title}** ({act.type}, {act.duration_minutes} min)")
            lines.append(f"Materials: {', '.join(act.materials)}")
            lines.append("Instructions:")
            lines += [f"  {i+1}. {step}" for i, step in enumerate(act.instructions)]
            lines.append(f"Success criteria: {', '.join(act.success_criteria)}")

        asmt = assessments_by_period.get(p.period_number)
        if asmt:
            lines.append("\n**Assessment**")
            for i, mcq in enumerate(asmt.mcqs, 1):
                lines.append(f"{i}. {mcq.question}")
                for j, opt in enumerate(mcq.options):
                    marker = "✅" if j == mcq.correct_index else "  "
                    lines.append(f"   {marker} {chr(97+j)}) {opt}")
            for wq in asmt.written_questions:
                lines.append(f"- ({wq.type}) {wq.question}")
                lines.append(f"  *Answer key:* {wq.answer_key}")
                lines.append(f"  *Rubric:* {wq.rubric}")

    if pkg.learning_gaps:
        lines.append("\n---\n## Learning Gap Analysis")
        for g in pkg.learning_gaps:
            lines.append(f"- **{g.misconception}** (severity: {g.severity})")
            lines.append(f"  - Diagnostic question: {g.diagnostic_question}")
            lines.append(f"  - Remedial action: {g.remedial_action}")

    v = pkg.validation
    lines.append("\n---\n## Validation Summary")
    lines.append(f"- Passed: {v.passed}")
    lines.append(f"- Hallucination risk score: {v.hallucination_score} (0=low, 1=high)")
    if v.issues:
        lines.append("- Issues:")
        lines += [f"  - [{i.severity}] ({i.stage}) {i.message}" for i in v.issues]

    return "\n".join(lines)


def render_html(markdown_text: str) -> str:
    body = md_lib.markdown(markdown_text, extensions=["tables", "fenced_code"])
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: -apple-system, Segoe UI, Roboto, sans-serif; max-width: 860px;
        margin: 2rem auto; padding: 0 1rem; line-height: 1.55; color: #1a1a1a; }}
h1, h2, h3 {{ color: #1d3557; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 2rem 0; }}
code {{ background: #f4f4f4; padding: 0 4px; border-radius: 3px; }}
</style></head><body>{body}</body></html>"""


def render_pdf(html_text: str) -> bytes:
    buf = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html_text), dest=buf)
    return buf.getvalue()


_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)")


def render_docx(markdown_text: str) -> bytes:
    """Minimal, dependency-free Markdown -> DOCX walker: headings and bullets
    map to Word styles; everything else becomes a plain paragraph. Good enough
    for a lesson-plan document without needing pandoc installed."""
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=min(level, 4))
            continue
        if line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\s*\d+\.\s+", "", line), style="List Number")
            continue
        clean = line.replace("**", "").replace("*", "").replace("---", "")
        if clean.strip():
            doc.add_paragraph(clean.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
